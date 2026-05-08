"""채팅 응답을 4 포맷 (DOCX/XLSX/CSV/TXT) 으로 변환하는 단일 책임 서비스.

옵션 B — Storage 백업 없음. 백엔드는 bytes 만 생성, Frontend 가 Blob 다운로드.

마크다운 변환은 단순 패서 (header / list / paragraph). 표/이미지는 Phase 4 비스코프.
"""

from __future__ import annotations

import csv
import io
import re
from typing import Iterable

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook


_HEADER_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_RE = re.compile(r"^[\-\*]\s+(.*)$")
_NUM_RE = re.compile(r"^\d+\.\s+(.*)$")
_TABLE_ROW_RE = re.compile(r"^\|(.+)\|$")


def _strip_inline_markdown(text: str) -> str:
    """**bold** / *italic* / `code` 마커를 단순 제거 (DOCX/CSV/TXT 평문화)."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"(?<!\*)\*(?!\s)(.+?)\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text


def _iter_lines(content: str) -> Iterable[str]:
    """공백 라인을 보존하며 한 줄씩 yield."""
    for raw in content.splitlines():
        yield raw.rstrip()


# ── DOCX ──────────────────────────────────────────────────────


def generate_docx(content: str, *, title: str = "AJIN AI 응답") -> bytes:
    """마크다운을 단순 DOCX 로 변환.

    헤더 ##, 리스트 -, 번호 1., 일반 단락만 처리.
    """
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)

    if title:
        h = doc.add_heading(title, level=1)
        h.alignment = 0  # left

    for line in _iter_lines(content):
        if not line.strip():
            doc.add_paragraph("")
            continue

        m_header = _HEADER_RE.match(line)
        if m_header:
            level = min(len(m_header.group(1)), 4)
            doc.add_heading(_strip_inline_markdown(m_header.group(2)), level=level)
            continue

        m_bullet = _BULLET_RE.match(line)
        if m_bullet:
            doc.add_paragraph(
                _strip_inline_markdown(m_bullet.group(1)),
                style="List Bullet",
            )
            continue

        m_num = _NUM_RE.match(line)
        if m_num:
            doc.add_paragraph(
                _strip_inline_markdown(m_num.group(1)),
                style="List Number",
            )
            continue

        doc.add_paragraph(_strip_inline_markdown(line))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── XLSX ──────────────────────────────────────────────────────


def generate_xlsx(content: str, *, sheet_name: str = "AI 응답") -> bytes:
    """마크다운을 XLSX 로 변환.

    | A | B | 형태의 GFM 표를 만나면 시트 행으로 직접 변환.
    그 외 라인은 A 열에 평문화하여 누적.
    """
    wb = Workbook()
    ws = wb.active
    if ws is None:
        ws = wb.create_sheet(sheet_name)
    else:
        ws.title = sheet_name[:31]  # XLSX sheet name limit

    row_idx = 1
    in_table = False
    table_separator_seen = False

    for line in _iter_lines(content):
        m_table = _TABLE_ROW_RE.match(line.strip())

        if m_table:
            cells_raw = m_table.group(1).split("|")
            cells = [_strip_inline_markdown(c.strip()) for c in cells_raw]

            # GFM 의 |---|---| 구분선은 데이터로 취급하지 않음
            if all(re.match(r"^:?-+:?$", c) for c in cells if c):
                table_separator_seen = True
                continue

            for col_idx, cell in enumerate(cells, start=1):
                ws.cell(row=row_idx, column=col_idx, value=cell)
            row_idx += 1
            in_table = True
            continue

        if in_table and not line.strip():
            in_table = False
            table_separator_seen = False
            row_idx += 1  # 빈 줄로 표 분리
            continue

        if not line.strip():
            row_idx += 1
            continue

        m_header = _HEADER_RE.match(line)
        if m_header:
            cell = ws.cell(row=row_idx, column=1, value=_strip_inline_markdown(m_header.group(2)))
            cell.font = cell.font.copy(bold=True)
            row_idx += 1
            continue

        m_bullet = _BULLET_RE.match(line) or _NUM_RE.match(line)
        if m_bullet:
            ws.cell(row=row_idx, column=1, value=f"• {_strip_inline_markdown(m_bullet.group(1))}")
            row_idx += 1
            continue

        ws.cell(row=row_idx, column=1, value=_strip_inline_markdown(line))
        row_idx += 1

    # 첫 컬럼 폭만 자동 확장
    if ws.column_dimensions:
        ws.column_dimensions["A"].width = 60

    buf = io.BytesIO()
    wb.save(buf)
    _ = table_separator_seen  # 향후 헤더 강조용으로 유지
    return buf.getvalue()


# ── CSV (utf-8-sig BOM — Excel 한글 호환) ─────────────────────


def generate_csv(content: str) -> bytes:
    """마크다운을 CSV 로 변환. utf-8-sig BOM 포함 (Excel 한글 깨짐 방지)."""
    buf = io.StringIO()
    writer = csv.writer(buf, quoting=csv.QUOTE_MINIMAL)

    for line in _iter_lines(content):
        m_table = _TABLE_ROW_RE.match(line.strip())
        if m_table:
            cells_raw = m_table.group(1).split("|")
            cells = [_strip_inline_markdown(c.strip()) for c in cells_raw]
            if all(re.match(r"^:?-+:?$", c) for c in cells if c):
                continue
            writer.writerow(cells)
            continue

        if not line.strip():
            writer.writerow([])
            continue

        m_header = _HEADER_RE.match(line)
        if m_header:
            writer.writerow([f"## {_strip_inline_markdown(m_header.group(2))}"])
            continue

        m_bullet = _BULLET_RE.match(line) or _NUM_RE.match(line)
        if m_bullet:
            writer.writerow([f"• {_strip_inline_markdown(m_bullet.group(1))}"])
            continue

        writer.writerow([_strip_inline_markdown(line)])

    return ("﻿" + buf.getvalue()).encode("utf-8")


# ── TXT (utf-8) ───────────────────────────────────────────────


def generate_txt(content: str) -> bytes:
    """마크다운을 평문 TXT 로. 인라인 마커만 제거, 줄 구조는 유지."""
    out_lines: list[str] = []
    for line in _iter_lines(content):
        out_lines.append(_strip_inline_markdown(line))
    return ("\n".join(out_lines) + "\n").encode("utf-8")


# ── 통합 디스패처 ─────────────────────────────────────────────


_MIME_BY_FORMAT: dict[str, tuple[str, str]] = {
    "docx": (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".docx",
    ),
    "xlsx": (
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".xlsx",
    ),
    "csv": ("text/csv; charset=utf-8", ".csv"),
    "txt": ("text/plain; charset=utf-8", ".txt"),
}


def generate(content: str, fmt: str) -> tuple[bytes, str, str]:
    """포맷 디스패처 — (bytes, mime, ext) 반환."""
    if fmt == "docx":
        data = generate_docx(content)
    elif fmt == "xlsx":
        data = generate_xlsx(content)
    elif fmt == "csv":
        data = generate_csv(content)
    elif fmt == "txt":
        data = generate_txt(content)
    else:
        raise ValueError(f"지원하지 않는 형식: {fmt}")
    mime, ext = _MIME_BY_FORMAT[fmt]
    return data, mime, ext
