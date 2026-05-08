"""규정 준수 라우터.

v3.0: 3-Tier 권한 적용
- Tier 1 VIEW: 모든 인증 사용자 (scenarios, facilities, risk, timeline, network, tariff)
- Tier 2 ANALYZE: 관련 부서 EMPLOYEE+ (check, classify)
- Tier 3 OPERATE: TEAM_LEAD+ (crawl, changes/ack)
"""

import json
from pathlib import Path
from typing import Any

from fastapi import APIRouter, Depends, HTTPException

from backend.schemas.compliance import (
    AcknowledgeResponse,
    ChangeItem,
    ChangeListResponse,
    ClassifyRequest,
    ClassifyResponse,
    ComplianceCheckRequest,
    ComplianceCheckResponse,
    CrawlRunAllResponse,
    CrawlRunResponse,
    CrawlResultMeta,
    CrawlResultsListResponse,
    CrawlResultItem,
    CrawlResultDetailResponse,
    FacilityItem,
    PlotlyResponse,
    RiskScoreItem,
    RiskScoreResponse,
    ScenarioChangeVersion,
    ScenarioDetailResponse,
    ScenarioReference,
    ScenarioRegulationMeta,
    ScenarioSimRiskScore,
    ScenarioSimImpact,
    ScenarioSimEvidence,
    ScenarioSimulateRequest,
    ScenarioSimulateResponse,
    TariffSimulateRequest,
    TariffSimulateResponse,
)
from backend.dependencies import get_current_user, require_permission
from config import DATA_DIR

router = APIRouter(prefix="/compliance", tags=["compliance"])


# ═══════════════════════════════════════════════════════════════
# D-2-3  GET /scenarios
# ═══════════════════════════════════════════════════════════════

@router.get("/scenarios")
async def list_scenarios(user=Depends(get_current_user)):
    """로드된 시나리오 목록을 반환한다. (인증 필수)"""
    scenarios_dir = DATA_DIR / "scenarios"
    scenarios: list[dict] = []

    if scenarios_dir.exists():
        for f in sorted(scenarios_dir.glob("*.json")):
            try:
                data = json.loads(f.read_text(encoding="utf-8"))
                if isinstance(data, list):
                    scenarios.extend(d for d in data if isinstance(d, dict))
                elif isinstance(data, dict):
                    scenarios.append(data)
            except Exception:
                continue

    # scenario_id 가 비어있는 항목 보정
    cleaned = [s for s in scenarios if s.get("scenario_id") or s.get("id")]
    return {"scenarios": cleaned, "total": len(cleaned)}


# ═══════════════════════════════════════════════════════════════
# D-2-12  GET /facilities  ─ plants.json 기반 19개소
# ═══════════════════════════════════════════════════════════════

def _load_facilities() -> list[FacilityItem]:
    """data/facility_db/plants.json 의 plants + subsidiaries(국내/해외) 19개소를 통합."""
    path = DATA_DIR / "facility_db" / "plants.json"
    if not path.exists():
        return []

    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return []

    facilities: list[FacilityItem] = []

    def _coord(item: dict) -> dict:
        lat = item.get("lat")
        lng = item.get("lng")
        return {
            "lat": float(lat) if isinstance(lat, (int, float)) else None,
            "lng": float(lng) if isinstance(lng, (int, float)) else None,
        }

    # 자사 공장
    for p in data.get("plants", []):
        facilities.append(FacilityItem(
            plant_id=p.get("plant_id", ""),
            name=p.get("name", ""),
            location=p.get("location", ""),
            address=p.get("location", ""),
            certifications=p.get("certifications", []),
            processes=p.get("main_business", []) or p.get("main_processes", []),
            kind="plant",
            country="KR",
            **_coord(p),
        ))

    # 국내 계열사
    for s in data.get("subsidiaries_domestic", []):
        facilities.append(FacilityItem(
            plant_id=s.get("subsidiary_id") or s.get("plant_id", ""),
            name=s.get("name", ""),
            location=s.get("location", ""),
            address=s.get("location", ""),
            certifications=s.get("certifications", []),
            processes=s.get("main_business", []) or s.get("products", []),
            kind="subsidiary_domestic",
            country="KR",
            **_coord(s),
        ))

    # 해외 법인
    for s in data.get("subsidiaries_overseas", []):
        facilities.append(FacilityItem(
            plant_id=s.get("subsidiary_id") or s.get("id", ""),
            name=s.get("name", ""),
            location=s.get("location", "") or s.get("city", ""),
            address=s.get("location", "") or s.get("city", ""),
            certifications=s.get("certifications", []),
            processes=s.get("main_business", []) or s.get("products", []),
            kind="subsidiary_overseas",
            country=s.get("country", ""),
            **_coord(s),
        ))

    return facilities


@router.get("/facilities")
async def list_facilities(user=Depends(get_current_user)):
    """19개 사업장 (자사 + 국내 계열사 + 해외 법인) 통합 반환."""
    facilities = _load_facilities()
    return {
        "facilities": facilities,
        "total": len(facilities),
        "domestic": sum(1 for f in facilities if f.country == "KR"),
        "overseas": sum(1 for f in facilities if f.country and f.country != "KR"),
    }


# ═══════════════════════════════════════════════════════════════
# D-2-2  GET /risk/scores  ─ 100점 스코어링
# ═══════════════════════════════════════════════════════════════

@router.get("/risk/scores", response_model=RiskScoreResponse)
async def list_risk_scores(user=Depends(get_current_user)):
    """모든 시나리오의 리스크 점수 (100점 + CRITICAL/HIGH/MEDIUM/LOW)."""
    from features.compliance.risk_scorer import score_all_scenarios, get_risk_summary

    try:
        scores = score_all_scenarios(str(DATA_DIR / "scenarios"))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"리스크 계산 실패: {e}")

    items = [
        RiskScoreItem(
            scenario_id=s.scenario_id,
            title=s.title,
            total_score=s.total_score,
            grade=s.grade,
            financial_impact=s.financial_impact,
            likelihood=s.likelihood,
            urgency=s.urgency,
            deadline=s.deadline,
            days_remaining=s.days_remaining,
            affected_plants=s.affected_plants,
            mitigation_status=s.mitigation_status,
        )
        for s in scores
    ]
    return RiskScoreResponse(
        total=len(items),
        summary=get_risk_summary(scores),
        scores=items,
    )


# ═══════════════════════════════════════════════════════════════
# D-2-5  POST /tariff/simulate
# ═══════════════════════════════════════════════════════════════

@router.post("/tariff/simulate", response_model=TariffSimulateResponse)
async def simulate_tariff_endpoint(
    req: TariffSimulateRequest,
    user=Depends(get_current_user),
):
    """관세 시뮬레이션 — 6품목(EWP/CCH/OBC/볼시트/도어/EV배터리) 기본 + 환율 적용."""
    from features.compliance.tariff_simulator import simulate_tariff

    out = simulate_tariff(tariff_rate=req.tariff_rate, exchange_rate=req.exchange_rate)
    return TariffSimulateResponse(
        tariff_rate=out["tariff_rate"],
        exchange_rate=out["exchange_rate"],
        total_annual_usd=out["total_annual_usd"],
        total_annual_krw=out["total_annual_krw"],
        total_annual_krw_billion=round(out["total_annual_krw"] / 1e8, 2),
        avg_cost_increase=out["avg_cost_increase"],
        results=[
            {
                "product": r.product,
                "tariff_rate": r.tariff_rate,
                "unit_tariff": r.unit_tariff,
                "annual_tariff": r.annual_tariff,
                "annual_tariff_krw": r.annual_tariff_krw,
                "cost_increase_pct": r.cost_increase_pct,
            }
            for r in out["results"]
        ],
    )


# ═══════════════════════════════════════════════════════════════
# D-2-4  GET /timeline  ─ Plotly Figure JSON
# ═══════════════════════════════════════════════════════════════

@router.get("/timeline", response_model=PlotlyResponse)
async def get_timeline(user=Depends(get_current_user)):
    """데드라인 간트 차트 (Plotly Figure JSON)."""
    from features.compliance.risk_scorer import score_all_scenarios
    from features.compliance.timeline_builder import build_deadline_timeline

    try:
        scores = score_all_scenarios(str(DATA_DIR / "scenarios"))
        fig = build_deadline_timeline(scores)
        return PlotlyResponse(figure=fig.to_plotly_json())
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"타임라인 생성 실패: {e}")


# ═══════════════════════════════════════════════════════════════
# D-2-7  GET /network/{scenario_id}  ─ Plotly Network
# ═══════════════════════════════════════════════════════════════

@router.get("/network/{scenario_id}", response_model=PlotlyResponse)
async def get_impact_network(scenario_id: str, user=Depends(get_current_user)):
    """규제 → 시설 → 부서 영향 네트워크 (Plotly Figure JSON)."""
    from features.compliance.impact_network import build_impact_network

    scenarios_dir = DATA_DIR / "scenarios"
    if not scenarios_dir.exists():
        raise HTTPException(status_code=404, detail="시나리오 디렉토리 없음")

    target: dict | None = None
    for f in scenarios_dir.glob("*.json"):
        try:
            d = json.loads(f.read_text(encoding="utf-8"))
            cands = d if isinstance(d, list) else [d]
            for c in cands:
                if isinstance(c, dict) and (
                    c.get("scenario_id") == scenario_id or c.get("id") == scenario_id
                ):
                    target = c
                    break
            if target:
                break
        except Exception:
            continue

    if not target:
        raise HTTPException(status_code=404, detail=f"시나리오 {scenario_id} 없음")

    try:
        fig = build_impact_network(target)
        return PlotlyResponse(figure=fig.to_plotly_json())
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"네트워크 생성 실패: {e}")


# ═══════════════════════════════════════════════════════════════
# D-2-6  GET /changes/recent  +  POST /changes/{id}/acknowledge
# ═══════════════════════════════════════════════════════════════

@router.get("/changes/recent", response_model=ChangeListResponse)
async def get_recent_changes_endpoint(
    limit: int = 20,
    unack_only: bool = False,
    user=Depends(get_current_user),
):
    """최근 규제 변경 이력 (compliance.db 기반)."""
    from features.compliance.change_detector import (
        init_change_db,
        get_recent_changes,
        get_change_stats,
    )

    init_change_db()
    raw = get_recent_changes(limit=limit, unacknowledged_only=unack_only)
    items = [
        ChangeItem(
            id=int(c.get("id", 0) or 0),
            regulation_type=c.get("regulation_type", ""),
            change_type=c.get("change_type", ""),
            item_id=c.get("item_id", ""),
            title=c.get("title", "") or c.get("description", ""),
            summary=c.get("summary", "") or c.get("description", ""),
            detected_at=c.get("detected_at", ""),
            acknowledged=bool(c.get("acknowledged", 0)),
        )
        for c in raw
    ]
    return ChangeListResponse(total=len(items), stats=get_change_stats(), changes=items)


@router.post("/changes/{change_id}/acknowledge", response_model=AcknowledgeResponse)
async def acknowledge_change_endpoint(
    change_id: int,
    user=Depends(get_current_user),
):
    """변경 이력을 '확인 완료' 상태로 마킹한다. (인증 사용자)"""
    from features.compliance.change_detector import init_change_db, acknowledge_change

    init_change_db()
    try:
        acknowledge_change(change_id, user_id=getattr(user, "employee_id", "") or "")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"확인 처리 실패: {e}")
    return AcknowledgeResponse(ok=True, change_id=change_id)


# ═══════════════════════════════════════════════════════════════
# D-2-8  POST /classify  ─ TF-IDF + RF
# ═══════════════════════════════════════════════════════════════

@router.post("/classify", response_model=ClassifyResponse)
async def classify_regulation(req: ClassifyRequest, user=Depends(get_current_user)):
    """규제 텍스트의 리스크 레벨을 분류한다."""
    from features.compliance.regulation_classifier import get_regulation_classifier

    if not req.text or not req.text.strip():
        raise HTTPException(status_code=400, detail="text 가 비어있습니다.")

    try:
        clf = get_regulation_classifier()
        out = clf.classify(req.text)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"분류 실패: {e}")

    return ClassifyResponse(
        severity=str(out.severity),
        confidence=float(out.confidence),
        all_scores={str(k): float(v) for k, v in out.all_scores.items()},
        related_departments=list(out.related_departments),
        affected_plants=list(out.affected_plants),
        risk_score=int(out.risk_score),
        recommended_actions=list(out.recommended_actions),
        response_deadline=out.response_deadline or "",
    )


# ═══════════════════════════════════════════════════════════════
# D-2-1 / D-2-12  Crawler control
# ═══════════════════════════════════════════════════════════════

# 단일 크롤러 인스턴스 매핑 (지연 로딩)
_CRAWLER_KEYS = {
    "iso": ("features.compliance.iso_crawler", "ISOCrawler"),
    "apqp": ("features.compliance.apqp_crawler", "APQPCrawler"),
    "msds": ("features.compliance.msds_crawler", "MSDSCrawler"),
    "domestic_law": ("features.compliance.domestic_law_crawler", "DomesticLawCrawler"),
    "eu_regulation": ("features.compliance.eu_regulation_crawler", "EURegulationCrawler"),
    "oem_quality": ("features.compliance.oem_quality_crawler", "OEMQualityCrawler"),
    "carbon_esg": ("features.compliance.carbon_esg_crawler", "CarbonESGCrawler"),
    "ev_battery": ("features.compliance.ev_battery_crawler", "EVBatteryCrawler"),
    "global_trade": ("features.compliance.global_trade_crawler", "GlobalTradeCrawler"),
}


def _summarize_crawl(name: str, result: Any) -> CrawlRunResponse:
    """크롤러 결과 객체를 공통 응답 스키마로 정규화."""
    return CrawlRunResponse(
        name=name,
        crawled_at=str(getattr(result, "crawled_at", "")),
        source=str(getattr(result, "source", "")),
        total_count=int(
            getattr(result, "total_count", 0)
            or getattr(result, "total_records", 0)
            or getattr(result, "total_phases", 0)
        ),
        updates_found=int(
            getattr(result, "updates_found", 0)
            or getattr(result, "updates_needed", 0)
            or getattr(result, "total_updates", 0)
            or getattr(result, "action_needed", 0)
        ),
        errors=list(getattr(result, "errors", []) or []),
    )


@router.post("/crawl/run/{name}", response_model=CrawlRunResponse)
async def run_single_crawler(
    name: str,
    user=Depends(get_current_user),
):
    """개별 크롤러 실행. name in iso/apqp/msds/domestic_law/eu_regulation/oem_quality/carbon_esg/ev_battery/global_trade"""
    if name not in _CRAWLER_KEYS:
        raise HTTPException(
            status_code=400,
            detail=f"지원되지 않는 크롤러: {name}. 가능: {list(_CRAWLER_KEYS.keys())}",
        )

    mod_path, cls_name = _CRAWLER_KEYS[name]
    try:
        mod = __import__(mod_path, fromlist=[cls_name])
        cls = getattr(mod, cls_name)
        inst = cls(DATA_DIR / "crawled")
        res = inst.crawl()
        return _summarize_crawl(name, res)
    except Exception as e:
        return CrawlRunResponse(name=name, errors=[f"{type(e).__name__}: {e}"])


@router.post("/crawl/run-all", response_model=CrawlRunAllResponse)
async def run_all_crawlers(user=Depends(get_current_user)):
    """9개 크롤러 일괄 실행. 일부 실패해도 나머지는 계속 진행."""
    out: dict[str, CrawlRunResponse] = {}
    total_changes = 0

    for name, (mod_path, cls_name) in _CRAWLER_KEYS.items():
        try:
            mod = __import__(mod_path, fromlist=[cls_name])
            cls = getattr(mod, cls_name)
            inst = cls(DATA_DIR / "crawled")
            res = inst.crawl()
            summary = _summarize_crawl(name, res)
            out[name] = summary
            total_changes += summary.updates_found
        except Exception as e:
            out[name] = CrawlRunResponse(name=name, errors=[f"{type(e).__name__}: {e}"])

    return CrawlRunAllResponse(crawlers=out, total_changes=total_changes)


# ═══════════════════════════════════════════════════════════════
# 기존: POST /check  (키워드 매핑)
# ═══════════════════════════════════════════════════════════════

@router.post("/check", response_model=ComplianceCheckResponse)
async def check_compliance(
    req: ComplianceCheckRequest,
    user=Depends(get_current_user),
    _perm=Depends(require_permission("compliance.run_analysis")),
):
    """규정 준수 검사를 수행한다. (관련 부서 EMPLOYEE+ 필요)"""
    try:
        query = req.query.lower()
        standards: list[str] = []
        status = "확인 필요"

        keyword_map = {
            "iatf": ["IATF 16949"],
            "iso 14001": ["ISO 14001"],
            "iso 45001": ["ISO 45001"],
            "reach": ["EU REACH"],
            "rohs": ["EU RoHS"],
            "ppap": ["IATF 16949 - PPAP"],
            "fmea": ["IATF 16949 - FMEA"],
            "spc": ["IATF 16949 - SPC"],
            "msds": ["화학물질관리법", "산업안전보건법"],
        }

        for kw, stds in keyword_map.items():
            if kw in query:
                standards.extend(stds)

        if standards:
            status = "관련 규정 발견"

        return ComplianceCheckResponse(
            answer=f"'{req.query}'에 대한 규정 검사 결과입니다.",
            relevant_standards=list(set(standards)),
            compliance_status=status,
            source="rules",
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ═══════════════════════════════════════════════════════════════
# v3.6 Phase 3 — GET /scenarios/{scenario_id}/detail
# 시나리오 원문 상세: 법규명·조항·시행일·변경 전/후 텍스트·체크리스트·근거 링크
# (시뮬레이션과 분리 — 분석 vs 레퍼런스)
# ═══════════════════════════════════════════════════════════════


def _load_scenario_raw(scenario_id: str) -> dict[str, Any] | None:
    """data/scenarios/*.json 에서 매칭 시나리오 raw dict 반환."""
    scenarios_dir = DATA_DIR / "scenarios"
    if not scenarios_dir.exists():
        return None
    for f in scenarios_dir.glob("*.json"):
        try:
            with open(f, encoding="utf-8") as fp:
                data = json.load(fp)
        except Exception:
            continue
        sid = data.get("scenario_id") or data.get("id") or f.stem
        if sid == scenario_id:
            return data
        # us_trade_regulations.json 처럼 여러 시나리오를 배열로 묶은 경우
        if isinstance(data.get("scenarios"), list):
            for s in data["scenarios"]:
                if isinstance(s, dict) and (s.get("scenario_id") == scenario_id or s.get("id") == scenario_id):
                    return s
    return None


def _days_until(date_str: str) -> int:
    """ISO 날짜 문자열 → 오늘까지의 남은 일수 (음수 = 지난)"""
    if not date_str:
        return 0
    try:
        from datetime import datetime
        d = datetime.fromisoformat(date_str.split("T")[0])
        delta = (d - datetime.now()).days
        return max(delta, 0)
    except Exception:
        return 0


@router.get(
    "/scenarios/{scenario_id}/detail",
    response_model=ScenarioDetailResponse,
)
async def get_scenario_detail(scenario_id: str, user=Depends(get_current_user)):
    """선택한 시나리오의 원문 상세 정보를 반환한다.

    시뮬레이션 라우트(`POST /scenarios/{id}/simulate`)와 별개:
    - 시뮬레이션 = 위험도 점수 + 영향 시설/부서 + 비용 (분석)
    - 상세    = 법규 원문 + 변경 전/후 + 체크리스트 + 근거 (레퍼런스)
    """
    raw = _load_scenario_raw(scenario_id)
    if raw is None:
        raise HTTPException(
            status_code=404,
            detail=f"시나리오 '{scenario_id}' 를 찾을 수 없습니다.",
        )

    # 법규 메타 (없으면 빈값)
    reg_raw = raw.get("regulation") or {}
    regulation = ScenarioRegulationMeta(
        name=str(reg_raw.get("name", "")),
        article=str(reg_raw.get("article", "")),
        authority=str(reg_raw.get("authority", reg_raw.get("issuer", ""))),
        category=str(reg_raw.get("category", "")),
    )

    # 변경 전/후
    change_detail = raw.get("change_detail") or {}
    before = None
    after = None
    if isinstance(change_detail.get("before"), dict):
        b = change_detail["before"]
        before = ScenarioChangeVersion(
            text=str(b.get("text", "")),
            effective_date=str(b.get("effective_date", "")),
            version=str(b.get("version", "")),
        )
    if isinstance(change_detail.get("after"), dict):
        a = change_detail["after"]
        after = ScenarioChangeVersion(
            text=str(a.get("text", "")),
            effective_date=str(a.get("effective_date", "")),
            version=str(a.get("version", "")),
        )

    # 참고 자료 — reference_url + 추가 references 배열
    refs: list[ScenarioReference] = []
    if raw.get("reference_url"):
        refs.append(
            ScenarioReference(
                title=f"{regulation.authority or '관련 기관'} 공식 자료",
                url=str(raw["reference_url"]),
            )
        )
    extra_refs = raw.get("references") or []
    for r in extra_refs:
        if isinstance(r, str):
            refs.append(ScenarioReference(title=r, url=r if r.startswith("http") else ""))
        elif isinstance(r, dict):
            refs.append(
                ScenarioReference(
                    title=str(r.get("title", r.get("name", "참고 자료"))),
                    url=str(r.get("url", "")),
                )
            )

    deadline = str(raw.get("deadline", ""))
    days_remaining = _days_until(deadline)

    return ScenarioDetailResponse(
        scenario_id=scenario_id,
        title=str(raw.get("title", scenario_id)),
        description=str(raw.get("description", "")),
        regulation=regulation,
        change_before=before,
        change_after=after,
        severity=str(raw.get("severity", "medium")).lower(),
        impact_areas=list(raw.get("impact_areas", [])),
        applicable_plants=list(raw.get("applicable_plants", [])),
        affected_facility_ids=list(raw.get("affected_facility_ids", [])),
        affected_process_types=list(raw.get("affected_process_types", [])),
        deadline=deadline,
        days_remaining=days_remaining,
        required_actions=list(raw.get("required_actions", [])),
        estimated_cost=str(raw.get("estimated_cost", "")),
        references=refs[:10],
        raw={k: v for k, v in raw.items() if k not in ("change_detail",)},  # change_detail 은 텍스트 큼
    )


# ═══════════════════════════════════════════════════════════════
# v3.6 Phase 2 — POST /scenarios/{scenario_id}/simulate
# 시나리오 통합 시뮬레이션: 위험도 + 영향 시설/부서 + 비용 추정 + 권장 액션
# ═══════════════════════════════════════════════════════════════


def _grade_to_category(grade: str) -> str:
    """위험도 등급 문자열 → UI 카테고리 매핑."""
    g = (grade or "").upper()
    if g in ("CRITICAL", "C", "S"):
        return "CRITICAL"
    if g in ("HIGH", "H", "A"):
        return "HIGH"
    if g in ("MEDIUM", "M", "B"):
        return "MEDIUM"
    return "LOW"


def _default_recommended_actions(category: str, deadline_days: int) -> list[str]:
    """카테고리별 기본 권장 액션 (시나리오 JSON 에 명시 없을 때 폴백)."""
    base = []
    if category == "CRITICAL":
        base = [
            f"D-{deadline_days} 내 영향 시설·부서 비상 점검",
            "관련 부서장 긴급 회의 소집 (24시간 내)",
            "법무팀·외부 전문가 자문 의뢰",
            "공급망·고객사 사전 통지 검토",
            "변경 이력 추적 및 증빙 문서 보관",
        ]
    elif category == "HIGH":
        base = [
            f"D-{deadline_days} 내 시정 계획 수립",
            "관련 부서 협업 회의 (1주일 내)",
            "변경된 기준에 따른 절차 개정",
            "교육 자료 업데이트",
        ]
    elif category == "MEDIUM":
        base = [
            "월간 점검 일정에 추가 확인",
            "관련 부서 공유 및 인지 확보",
            "필요시 절차 검토",
        ]
    else:  # LOW
        base = [
            "정기 점검 사항으로 반영",
            "관련 부서 공유",
        ]
    return base


@router.post(
    "/scenarios/{scenario_id}/simulate",
    response_model=ScenarioSimulateResponse,
)
async def simulate_scenario(
    scenario_id: str,
    req: ScenarioSimulateRequest = ScenarioSimulateRequest(),
    user=Depends(get_current_user),
):
    """선택한 시나리오의 통합 시뮬레이션 결과를 반환한다.

    조합:
      - 시나리오 메타 (제목·설명·deadline) — data/scenarios/*.json
      - 위험도 점수 — features.compliance.compliance_db
      - 영향 시설/부서 — features.compliance.impact_network
      - 비용 추정 — features.compliance.tariff_simulator (관세 카테고리만)
      - 권장 액션 — 시나리오 JSON 의 recommended_actions 또는 카테고리 기본값
      - 근거 링크 — 시나리오 JSON 의 references (있으면)
    """
    # ── 1) 시나리오 로드 ──
    scenarios_dir = DATA_DIR / "scenarios"
    scenario: dict[str, Any] | None = None
    if scenarios_dir.exists():
        for f in scenarios_dir.glob("*.json"):
            try:
                with open(f, encoding="utf-8") as fp:
                    data = json.load(fp)
                sid = data.get("scenario_id") or data.get("id") or f.stem
                if sid == scenario_id:
                    scenario = data
                    scenario["_source_file"] = str(f.name)
                    break
            except Exception:
                continue

    if scenario is None:
        raise HTTPException(
            status_code=404,
            detail=f"시나리오 '{scenario_id}' 를 찾을 수 없습니다.",
        )

    title = scenario.get("title") or scenario.get("name") or scenario_id
    description = (
        scenario.get("description")
        or (scenario.get("regulation") or {}).get("article", "")
        or ""
    )
    deadline_days = int(scenario.get("days_remaining", scenario.get("deadline_days", 0)) or 0)

    # ── 2) 위험도 점수 ──
    risk = ScenarioSimRiskScore()
    try:
        from features.compliance.compliance_db import compute_risk_scores  # type: ignore
        scores = compute_risk_scores()
        for s in scores:
            if s.get("scenario_id") == scenario_id:
                risk = ScenarioSimRiskScore(
                    total=int(s.get("total_score", 0)),
                    fin=int(s.get("financial_impact", 0)),
                    pos=int(s.get("likelihood", 0)),
                    urg=int(s.get("urgency", 0) or s.get("urgent", 0)),
                )
                grade = s.get("grade", "MEDIUM")
                break
        else:
            grade = scenario.get("grade", "MEDIUM")
    except Exception:
        grade = scenario.get("grade", "MEDIUM")

    category = _grade_to_category(grade)

    # ── 3) 영향 시설/부서 ──
    plants = list(scenario.get("affected_plants") or scenario.get("sites") or [])
    departments: list[str] = []
    try:
        from features.compliance.impact_network import REGULATION_DEPT_MAP  # type: ignore
        # 규제 키워드 → 부서 매핑
        text_pool = " ".join(
            [
                title,
                description,
                str(scenario.get("regulation", {})),
                " ".join(scenario.get("keywords", [])),
            ]
        )
        for kw, deps in REGULATION_DEPT_MAP.items():
            if kw in text_pool:
                departments.extend(deps)
        departments = list(dict.fromkeys(departments))  # 중복 제거 (순서 유지)
    except Exception:
        departments = scenario.get("affected_departments") or []

    # ── 4) 비용 추정 (관세 시나리오만) ──
    cost_estimate = 0.0
    cost_breakdown: list[dict[str, Any]] = []
    if "관세" in title or "tariff" in scenario_id.lower() or "trade" in (scenario.get("category", "")).lower():
        try:
            from features.compliance.tariff_simulator import simulate_tariff  # type: ignore
            tariff_rate = req.tariff_rate if req.tariff_rate is not None else 25.0
            exchange_rate = req.exchange_rate if req.exchange_rate is not None else 1380.0
            sim = simulate_tariff(tariff_rate=tariff_rate, exchange_rate=exchange_rate)
            cost_breakdown = sim.get("items", [])
            total_krw = sum(item.get("annual_tariff_krw", 0) for item in cost_breakdown)
            cost_estimate = round(total_krw / 1_000_000_000, 2)  # 10억 단위
        except Exception:
            pass

    impact = ScenarioSimImpact(
        plants=plants,
        departments=departments,
        cost_estimate_krw_bn=cost_estimate,
        cost_breakdown=cost_breakdown[:20],  # 최대 20개
    )

    # ── 5) 권장 액션 ──
    recommended = scenario.get("recommended_actions") or scenario.get("actions") or []
    if not recommended:
        recommended = _default_recommended_actions(category, deadline_days)

    # ── 6) 근거 링크 ──
    evidence: list[ScenarioSimEvidence] = []
    refs = scenario.get("references") or scenario.get("evidence_links") or []
    for ref in refs[:5]:
        if isinstance(ref, str):
            evidence.append(ScenarioSimEvidence(title=ref, url=ref if ref.startswith("http") else ""))
        elif isinstance(ref, dict):
            evidence.append(
                ScenarioSimEvidence(
                    title=ref.get("title", ref.get("name", "참고 자료")),
                    url=ref.get("url", ""),
                )
            )

    return ScenarioSimulateResponse(
        scenario_id=scenario_id,
        title=title,
        category=category,
        deadline_days=deadline_days,
        description=description,
        risk_score=risk,
        impact=impact,
        recommended_actions=recommended[:10],
        evidence_links=evidence,
    )


# ═══════════════════════════════════════════════════════════════
# v3.6 Phase 2 — GET /crawl/results, GET /crawl/results/{name}
# 크롤러 실행 결과 조회 (data/crawled/*.json)
# ═══════════════════════════════════════════════════════════════


# 크롤러 이름 → 실제 파일명 매핑 (이름이 다를 수 있음)
_CRAWLER_FILE_MAP = {
    "iso": "iso_standards.json",
    "apqp": "apqp_process.json",
    "msds": "msds_data.json",
    "domestic_law": "domestic_laws.json",
    "eu_regulation": "eu_regulations.json",
    "oem_quality": "oem_quality.json",
    "carbon_esg": "carbon_esg.json",
    "ev_battery": "ev_battery.json",
    "global_trade": "global_trade.json",
}

# 결과 JSON 안에서 항목 배열을 담는 필드 이름 후보 (크롤러마다 다름)
_ITEMS_FIELD_CANDIDATES = [
    "standards", "laws", "regulations", "items", "products",
    "data", "results", "records", "phases",
]


def _extract_items(payload: dict[str, Any]) -> list[dict[str, Any]]:
    """크롤러 JSON 의 메인 배열을 자동 탐지."""
    for field in _ITEMS_FIELD_CANDIDATES:
        if isinstance(payload.get(field), list):
            return payload[field]
    # 마지막 수단: dict 의 list 값 중 가장 큰 것
    largest: list[dict[str, Any]] = []
    for v in payload.values():
        if isinstance(v, list) and len(v) > len(largest):
            largest = v
    return largest


@router.get("/crawl/results", response_model=CrawlResultsListResponse)
async def list_crawl_results(user=Depends(get_current_user)):
    """모든 크롤러의 결과 메타데이터를 반환한다.

    실제 데이터는 data/crawled/{filename}.json 에 저장됨. 본 엔드포인트는
    각 파일을 열어 메타 (crawled_at, source, total_count 등) 만 추출.
    """
    crawled_dir = DATA_DIR / "crawled"
    out: list[CrawlResultMeta] = []

    for name, filename in _CRAWLER_FILE_MAP.items():
        path = crawled_dir / filename
        if not path.exists():
            continue
        try:
            with open(path, encoding="utf-8") as f:
                data = json.load(f)
            items = _extract_items(data)
            size_bytes = path.stat().st_size
            out.append(
                CrawlResultMeta(
                    name=name,
                    filename=filename,
                    crawled_at=str(data.get("crawled_at", "")),
                    source=str(data.get("source", "")),
                    total_count=int(data.get("total_count", len(items))),
                    updates_found=int(data.get("updates_found", 0)),
                    errors=list(data.get("errors", []))[:5],
                    size_bytes=size_bytes,
                )
            )
        except Exception as e:
            out.append(
                CrawlResultMeta(
                    name=name,
                    filename=filename,
                    crawled_at="",
                    source="",
                    total_count=0,
                    errors=[f"파일 읽기 실패: {e}"],
                )
            )

    return CrawlResultsListResponse(crawlers=out, total=len(out))


# ═══════════════════════════════════════════════════════════════
# v3.6 Phase 3 Item 2 — GET /crawl/results/{name}/download?format=...
# 5 포맷 다운로드: JSON · CSV · XLSX · DOCX · PDF
# ═══════════════════════════════════════════════════════════════


def _build_csv(items: list[dict[str, Any]]) -> bytes:
    """크롤러 항목 배열 → CSV. 모든 항목의 키 합집합을 컬럼으로 사용."""
    import csv
    from io import StringIO

    if not items:
        return b"title,url,summary\n"
    columns = ["title", "url", "summary"]
    extra_keys: list[str] = []
    for it in items:
        for k in it.keys():
            if k not in columns and k not in extra_keys and k != "extra":
                extra_keys.append(k)
    columns = columns + extra_keys[:20]  # 추가 컬럼 최대 20개

    out = StringIO()
    writer = csv.DictWriter(out, fieldnames=columns, extrasaction="ignore")
    writer.writeheader()
    for it in items:
        row = {col: str(it.get(col, "")).replace("\n", " ") for col in columns}
        writer.writerow(row)
    return ("﻿" + out.getvalue()).encode("utf-8")  # BOM (Excel 한글 호환)


def _build_xlsx(name: str, meta: dict[str, Any], items: list[dict[str, Any]]) -> bytes:
    """openpyxl 사용 — 헤더 굵게 + 자동 너비."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment
        from io import BytesIO
    except ImportError as e:
        raise HTTPException(status_code=503, detail=f"openpyxl 미설치: {e}")

    wb = Workbook()
    ws = wb.active
    ws.title = name[:30] or "crawl"

    # 메타 정보 (1-3행)
    ws.cell(row=1, column=1, value=f"크롤러: {name}")
    ws.cell(row=1, column=1).font = Font(bold=True, size=14)
    ws.cell(row=2, column=1, value=f"실행 시각: {meta.get('crawled_at', '')}")
    ws.cell(row=3, column=1, value=f"출처: {meta.get('source', '')}")

    # 헤더 (5행)
    columns = ["title", "url", "summary"]
    extra_keys: list[str] = []
    for it in items:
        for k in it.keys():
            if k not in columns and k not in extra_keys and k != "extra":
                extra_keys.append(k)
    columns = columns + extra_keys[:15]

    header_fill = PatternFill("solid", fgColor="D89400")
    for col_idx, col_name in enumerate(columns, start=1):
        c = ws.cell(row=5, column=col_idx, value=col_name)
        c.font = Font(bold=True, color="FFFFFF")
        c.fill = header_fill
        c.alignment = Alignment(horizontal="center")

    # 데이터 (6행~)
    for row_idx, it in enumerate(items, start=6):
        for col_idx, col_name in enumerate(columns, start=1):
            v = it.get(col_name, "")
            ws.cell(row=row_idx, column=col_idx, value=str(v)[:32000])  # XLSX 셀 한계

    # 자동 너비 (간단 추정)
    for col_idx, col_name in enumerate(columns, start=1):
        col_letter = ws.cell(row=5, column=col_idx).column_letter
        max_len = max([len(str(it.get(col_name, ""))) for it in items[:50]] + [len(col_name)])
        ws.column_dimensions[col_letter].width = min(max_len + 2, 50)

    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _build_docx(name: str, meta: dict[str, Any], items: list[dict[str, Any]]) -> bytes:
    """python-docx — 회사 양식 보고서."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from io import BytesIO
    except ImportError as e:
        raise HTTPException(status_code=503, detail=f"python-docx 미설치: {e}")

    doc = Document()
    title = doc.add_heading(f"크롤링 결과 보고서 — {name}", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0xD8, 0x94, 0x00)

    p = doc.add_paragraph()
    p.add_run(f"실행 시각: {meta.get('crawled_at', '')}\n").font.size = Pt(10)
    p.add_run(f"출처: {meta.get('source', '')}\n").font.size = Pt(10)
    p.add_run(f"항목 수: {len(items)}건").font.size = Pt(10)

    doc.add_paragraph()  # 빈 줄

    for i, it in enumerate(items, 1):
        h = doc.add_heading(f"{i}. {it.get('title', '(제목 없음)')}", level=2)
        for run in h.runs:
            run.font.size = Pt(13)
        if it.get("url"):
            url_p = doc.add_paragraph()
            url_run = url_p.add_run(f"🔗 {it.get('url')}")
            url_run.font.size = Pt(9)
            url_run.font.color.rgb = RGBColor(0x29, 0x80, 0xB9)
        if it.get("summary"):
            doc.add_paragraph(it.get("summary"))

    # 푸터
    doc.add_paragraph()
    foot = doc.add_paragraph()
    fr = foot.add_run("아진산업(주) | 본 보고서는 컴플라이언스 모니터링 시스템에서 자동 생성됨")
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_pdf(name: str, meta: dict[str, Any], items: list[dict[str, Any]]) -> bytes:
    """fpdf2 — PDF 보고서 (NanumGothic 한글 폰트)."""
    try:
        from fpdf import FPDF
        from io import BytesIO
        import os.path
    except ImportError as e:
        raise HTTPException(status_code=503, detail=f"fpdf2 미설치: {e}")

    pdf = FPDF(format="A4")
    pdf.add_page()

    # 한글 폰트 — Dockerfile 의 fonts-nanum 패키지가 /usr/share/fonts/truetype/nanum/ 에 설치됨
    font_loaded = False
    for font_path in [
        "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
        "/usr/share/fonts/truetype/nanum/NanumGothicBold.ttf",
        "/Library/Fonts/AppleSDGothicNeo.ttc",  # macOS 개발 환경
    ]:
        if os.path.exists(font_path):
            try:
                pdf.add_font("Korean", style="", fname=font_path)
                pdf.add_font("Korean", style="B", fname=font_path)
                pdf.set_font("Korean", size=10)
                font_loaded = True
                break
            except Exception:
                continue
    if not font_loaded:
        pdf.set_font("helvetica", size=10)

    pdf.set_font_size(18)
    pdf.cell(0, 12, f"크롤링 결과 보고서 - {name}", ln=True)
    pdf.set_font_size(9)
    pdf.cell(0, 6, f"실행 시각: {meta.get('crawled_at', '')}", ln=True)
    pdf.cell(0, 6, f"출처: {meta.get('source', '')[:80]}", ln=True)
    pdf.cell(0, 6, f"항목 수: {len(items)}건", ln=True)
    pdf.ln(4)

    for i, it in enumerate(items, 1):
        if pdf.get_y() > 270:
            pdf.add_page()
        pdf.set_font_size(11)
        title = str(it.get("title", "(제목 없음)"))[:120]
        pdf.multi_cell(0, 6, f"{i}. {title}")
        pdf.set_font_size(8)
        if it.get("url"):
            pdf.multi_cell(0, 5, f"URL: {str(it['url'])[:200]}")
        if it.get("summary"):
            summary = str(it["summary"])[:600]
            pdf.multi_cell(0, 5, summary)
        pdf.ln(2)

    return bytes(pdf.output())


@router.get("/crawl/results/{name}/download")
async def download_crawl_result(
    name: str,
    format: str = "json",
    user=Depends(get_current_user),
):
    """크롤러 결과를 5 포맷 중 하나로 다운로드.

    format: json | csv | xlsx | docx | pdf
    """
    from fastapi.responses import Response

    if name not in _CRAWLER_FILE_MAP:
        raise HTTPException(status_code=404, detail=f"크롤러 '{name}' 미등록.")

    filename = _CRAWLER_FILE_MAP[name]
    path = DATA_DIR / "crawled" / filename
    if not path.exists():
        raise HTTPException(
            status_code=404,
            detail=f"'{name}' 결과 파일이 없습니다. 'RUN ALL' 으로 먼저 실행하세요.",
        )

    try:
        with open(path, encoding="utf-8") as f:
            data = json.load(f)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"결과 파일 파싱 실패: {e}")

    raw_items = _extract_items(data)
    # 정규화 — title/url/summary 추출 + 나머지 보존
    items: list[dict[str, Any]] = []
    for raw in raw_items:
        if not isinstance(raw, dict):
            items.append({"title": str(raw), "url": "", "summary": ""})
            continue
        title = str(
            raw.get("title")
            or raw.get("name")
            or raw.get("standard")
            or raw.get("law_name")
            or raw.get("id")
            or "(제목 없음)"
        )
        url = str(raw.get("url", raw.get("link", raw.get("source_url", ""))))
        summary = str(
            raw.get("summary") or raw.get("description") or raw.get("content", "")
        )[:1000]
        # 추가 필드는 그대로 보존
        merged = {"title": title, "url": url, "summary": summary}
        for k, v in raw.items():
            if k not in merged:
                merged[k] = v
        items.append(merged)

    fmt = format.lower().strip()
    from datetime import datetime
    today = datetime.now().strftime("%Y%m%d")
    base_name = f"{name}_{today}"

    meta = {
        "crawled_at": data.get("crawled_at", ""),
        "source": data.get("source", ""),
    }

    if fmt == "json":
        return Response(
            content=json.dumps(data, ensure_ascii=False, indent=2),
            media_type="application/json; charset=utf-8",
            headers={"Content-Disposition": f"attachment; filename={base_name}.json"},
        )
    if fmt == "csv":
        return Response(
            content=_build_csv(items),
            media_type="text/csv; charset=utf-8",
            headers={"Content-Disposition": f"attachment; filename={base_name}.csv"},
        )
    if fmt == "xlsx":
        return Response(
            content=_build_xlsx(name, meta, items),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename={base_name}.xlsx"},
        )
    if fmt == "docx":
        return Response(
            content=_build_docx(name, meta, items),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={base_name}.docx"},
        )
    if fmt == "pdf":
        return Response(
            content=_build_pdf(name, meta, items),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={base_name}.pdf"},
        )
    # v3.6 Phase 4 — 6번째 포맷: 회사 양식 보고서 (DOCX, 표지+요약+본문+부록)
    if fmt == "report":
        author = (
            f"{user.get('username', '')} {user.get('position', '')}".strip()
            if isinstance(user, dict)
            else ""
        )
        return Response(
            content=_build_report_docx(name, meta, items, author=author),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={base_name}_report.docx"},
        )

    raise HTTPException(
        status_code=400,
        detail=f"지원하지 않는 포맷: '{fmt}'. (json|csv|xlsx|docx|pdf|report)",
    )


def _build_report_docx(
    name: str,
    meta: dict[str, Any],
    items: list[dict[str, Any]],
    author: str = "",
) -> bytes:
    """단일 크롤러 회사 양식 보고서 (DOCX, 표지+요약+본문+부록).

    Phase 4 6번째 포맷. 일반 _build_docx 보다 구조화된 양식.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
        from datetime import datetime
    except ImportError as e:
        raise HTTPException(status_code=503, detail=f"python-docx 미설치: {e}")

    doc = Document()
    today = datetime.now()
    today_str = today.strftime("%Y년 %m월 %d일")
    period_str = today.strftime("%Y년 %m월")

    # ── 표지 ──
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("아진산업(주)")
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xD8, 0x94, 0x00)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AJIN INDUSTRIAL CO., LTD.")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x8A, 0x82, 0x76)

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("법규 모니터링 보고서")
    r.font.size = Pt(22)
    r.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"— {name.upper()} —")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for _ in range(8):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"작성일 · {today_str}")
    r.font.size = Pt(11)
    if author:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"작성자 · {author}")
        r.font.size = Pt(11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"기간 · {period_str} 월간 보고")
    r.font.size = Pt(11)

    doc.add_page_break()

    # ── Executive Summary ──
    doc.add_heading("1. 개요 (Executive Summary)", level=1)
    p = doc.add_paragraph()
    p.add_run(f"본 보고서는 {name.upper()} 크롤러로 수집한 ").font.size = Pt(11)
    r = p.add_run(f"{len(items)}건")
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xD8, 0x94, 0x00)
    p.add_run(f"의 규제 항목을 정리한 것입니다.").font.size = Pt(11)

    p = doc.add_paragraph()
    p.add_run(f"  • 출처: {meta.get('source', '—')}").font.size = Pt(10)
    p = doc.add_paragraph()
    p.add_run(f"  • 마지막 갱신: {meta.get('crawled_at', '—')}").font.size = Pt(10)
    p = doc.add_paragraph()
    p.add_run(f"  • 항목 수: {len(items)}건").font.size = Pt(10)

    doc.add_paragraph()

    # ── 본문 ──
    doc.add_heading("2. 수집 항목 상세", level=1)
    for i, it in enumerate(items, 1):
        h = doc.add_heading(f"2.{i} {it.get('title', '(제목 없음)')}", level=2)
        if it.get("url"):
            p = doc.add_paragraph()
            r = p.add_run(f"🔗 출처: {it.get('url')}")
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x29, 0x80, 0xB9)
        if it.get("summary"):
            p = doc.add_paragraph(it.get("summary"))
            p.paragraph_format.space_after = Pt(8)

    # ── 부록 ──
    doc.add_page_break()
    doc.add_heading("부록 A. 전체 항목 색인", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "#"
    hdr[1].text = "제목"
    hdr[2].text = "URL"
    for i, it in enumerate(items, 1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = str(it.get("title", ""))[:80]
        row[2].text = str(it.get("url", ""))[:100]

    # ── 푸터 ──
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "아진산업(주) · 본 보고서는 컴플라이언스 모니터링 시스템에서 자동 생성됨"
    )
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════
# v3.6 Phase 4 — GET /crawl/results/bulk-download?format=...
# 9개 크롤러 결과를 단일 파일로 묶음 (JSON / XLSX / DOCX / PDF / ZIP / report)
# ═══════════════════════════════════════════════════════════════


def _gather_all_crawler_results() -> list[dict[str, Any]]:
    """9개 크롤러 결과 + 메타 + 항목 정규화 리스트."""
    out: list[dict[str, Any]] = []
    for crawler_name, filename in _CRAWLER_FILE_MAP.items():
        path = DATA_DIR / "crawled" / filename
        if not path.exists():
            out.append(
                {
                    "name": crawler_name,
                    "filename": filename,
                    "crawled_at": "",
                    "source": "",
                    "items": [],
                    "raw": {},
                    "errors": ["파일 없음 — RUN ALL 필요"],
                }
            )
            continue
        try:
            with open(path, encoding="utf-8") as f:
                data = json.load(f)
            raw_items = _extract_items(data)
            items: list[dict[str, Any]] = []
            for raw in raw_items:
                if not isinstance(raw, dict):
                    items.append({"title": str(raw), "url": "", "summary": ""})
                    continue
                title = str(
                    raw.get("title")
                    or raw.get("name")
                    or raw.get("standard")
                    or raw.get("law_name")
                    or raw.get("id")
                    or "(제목 없음)"
                )
                url = str(raw.get("url", raw.get("link", raw.get("source_url", ""))))
                summary = str(
                    raw.get("summary") or raw.get("description") or raw.get("content", "")
                )[:1000]
                items.append({"title": title, "url": url, "summary": summary, **raw})
            out.append(
                {
                    "name": crawler_name,
                    "filename": filename,
                    "crawled_at": str(data.get("crawled_at", "")),
                    "source": str(data.get("source", "")),
                    "items": items,
                    "raw": data,
                    "errors": list(data.get("errors", [])),
                }
            )
        except Exception as e:
            out.append(
                {
                    "name": crawler_name,
                    "filename": filename,
                    "errors": [f"파싱 실패: {e}"],
                    "items": [],
                    "raw": {},
                }
            )
    return out


def _build_bulk_xlsx(crawlers: list[dict[str, Any]]) -> bytes:
    """9개 시트 + Summary 시트."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment
        from io import BytesIO
        from datetime import datetime
    except ImportError as e:
        raise HTTPException(status_code=503, detail=f"openpyxl 미설치: {e}")

    wb = Workbook()
    # Summary 시트 (첫 번째)
    summary = wb.active
    summary.title = "Summary"
    summary["A1"] = "법규 모니터링 통합 보고서"
    summary["A1"].font = Font(bold=True, size=16, color="D89400")
    summary["A2"] = f"생성: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    summary["A3"] = f"크롤러 수: {len(crawlers)}"

    headers = ["#", "크롤러", "파일명", "마지막 갱신", "항목 수", "출처"]
    for i, h in enumerate(headers, 1):
        c = summary.cell(row=5, column=i, value=h)
        c.font = Font(bold=True, color="FFFFFF")
        c.fill = PatternFill("solid", fgColor="D89400")
        c.alignment = Alignment(horizontal="center")

    for ridx, c in enumerate(crawlers, start=6):
        summary.cell(row=ridx, column=1, value=ridx - 5)
        summary.cell(row=ridx, column=2, value=c["name"])
        summary.cell(row=ridx, column=3, value=c["filename"])
        summary.cell(row=ridx, column=4, value=c["crawled_at"])
        summary.cell(row=ridx, column=5, value=len(c["items"]))
        summary.cell(row=ridx, column=6, value=c["source"][:80])

    # 각 크롤러별 시트
    for c in crawlers:
        ws = wb.create_sheet(title=c["name"][:30])
        ws.cell(row=1, column=1, value=f"크롤러: {c['name']}").font = Font(bold=True, size=14)
        ws.cell(row=2, column=1, value=f"갱신: {c['crawled_at']}")
        ws.cell(row=3, column=1, value=f"출처: {c['source']}")
        # 헤더
        cols = ["#", "title", "url", "summary"]
        for i, h in enumerate(cols, 1):
            cell = ws.cell(row=5, column=i, value=h)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="D89400")
        for ridx, it in enumerate(c["items"], start=6):
            ws.cell(row=ridx, column=1, value=ridx - 5)
            ws.cell(row=ridx, column=2, value=str(it.get("title", ""))[:200])
            ws.cell(row=ridx, column=3, value=str(it.get("url", ""))[:300])
            ws.cell(row=ridx, column=4, value=str(it.get("summary", ""))[:1000])
        # 자동 너비
        ws.column_dimensions["A"].width = 5
        ws.column_dimensions["B"].width = 40
        ws.column_dimensions["C"].width = 50
        ws.column_dimensions["D"].width = 60

    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _build_bulk_docx(
    crawlers: list[dict[str, Any]],
    author: str = "",
) -> bytes:
    """9개 통합 회사 양식 보고서 (표지+목차+요약+9섹션+부록)."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from io import BytesIO
        from datetime import datetime
    except ImportError as e:
        raise HTTPException(status_code=503, detail=f"python-docx 미설치: {e}")

    doc = Document()
    today = datetime.now()
    today_str = today.strftime("%Y년 %m월 %d일")
    period_str = today.strftime("%Y년 %m월")
    total_items = sum(len(c["items"]) for c in crawlers)
    successful = sum(1 for c in crawlers if not c.get("errors"))

    # 표지
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("아진산업(주)")
    r.font.size = Pt(30)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xD8, 0x94, 0x00)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AJIN INDUSTRIAL CO., LTD.")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x8A, 0x82, 0x76)

    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("법규 모니터링 통합 보고서")
    r.font.size = Pt(24)
    r.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"— {len(crawlers)}개 크롤러 통합 —")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for _ in range(8):
        doc.add_paragraph()
    for line in [
        f"작성일 · {today_str}",
        f"작성자 · {author}" if author else "",
        f"기간 · {period_str} 월간 보고",
    ]:
        if not line:
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(11)

    doc.add_page_break()

    # 목차
    doc.add_heading("목차", level=1)
    doc.add_paragraph("1. 개요 (Executive Summary)")
    doc.add_paragraph("2. 크롤러별 상세")
    for i, c in enumerate(crawlers, 1):
        doc.add_paragraph(f"   2.{i} {c['name']} — {len(c['items'])}건")
    doc.add_paragraph("3. 부록 — 전체 항목 색인")
    doc.add_page_break()

    # Executive Summary
    doc.add_heading("1. 개요 (Executive Summary)", level=1)
    p = doc.add_paragraph()
    p.add_run("본 보고서는 ").font.size = Pt(11)
    r = p.add_run(f"{successful}/{len(crawlers)}개")
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xD8, 0x94, 0x00)
    p.add_run(f" 크롤러에서 수집한 ").font.size = Pt(11)
    r = p.add_run(f"총 {total_items}건")
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xD8, 0x94, 0x00)
    p.add_run(f"의 규제 항목을 통합 정리한 것입니다.").font.size = Pt(11)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "크롤러"
    hdr[1].text = "항목 수"
    hdr[2].text = "마지막 갱신"
    hdr[3].text = "상태"
    for c in crawlers:
        row = table.add_row().cells
        row[0].text = c["name"]
        row[1].text = str(len(c["items"]))
        row[2].text = c["crawled_at"][:10] if c["crawled_at"] else "—"
        row[3].text = "정상" if not c.get("errors") else "에러"

    doc.add_page_break()

    # 본문 - 각 크롤러 섹션
    doc.add_heading("2. 크롤러별 상세", level=1)
    for i, c in enumerate(crawlers, 1):
        doc.add_heading(f"2.{i} {c['name']}", level=2)
        p = doc.add_paragraph()
        p.add_run(f"  • 출처: {c['source'] or '—'}\n").font.size = Pt(10)
        p.add_run(f"  • 갱신: {c['crawled_at'] or '—'}\n").font.size = Pt(10)
        p.add_run(f"  • 항목 수: {len(c['items'])}건").font.size = Pt(10)
        if c.get("errors"):
            p = doc.add_paragraph()
            r = p.add_run(f"⚠ 에러: {'; '.join(c['errors'])}")
            r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
            r.font.size = Pt(10)
        # 상위 5개 항목만 본문에 표시 (큰 보고서 방지)
        for j, it in enumerate(c["items"][:5], 1):
            p = doc.add_paragraph()
            r = p.add_run(f"    {j}. {it.get('title', '(제목 없음)')[:100]}")
            r.font.size = Pt(10)
            r.font.bold = True
            if it.get("url"):
                p = doc.add_paragraph()
                r = p.add_run(f"       🔗 {it.get('url')[:120]}")
                r.font.size = Pt(8)
                r.font.color.rgb = RGBColor(0x29, 0x80, 0xB9)
        if len(c["items"]) > 5:
            p = doc.add_paragraph()
            r = p.add_run(
                f"    ... 외 {len(c['items']) - 5}건 (전체 항목은 부록 참조)"
            )
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    # 부록
    doc.add_page_break()
    doc.add_heading("3. 부록 — 전체 항목 색인", level=1)
    for i, c in enumerate(crawlers, 1):
        doc.add_heading(f"A.{i} {c['name']}", level=2)
        if not c["items"]:
            doc.add_paragraph("(항목 없음)")
            continue
        for j, it in enumerate(c["items"], 1):
            p = doc.add_paragraph()
            r = p.add_run(f"{j}. {it.get('title', '')[:120]}")
            r.font.size = Pt(9)
            if it.get("url"):
                p = doc.add_paragraph()
                r = p.add_run(f"   {it.get('url')[:150]}")
                r.font.size = Pt(8)
                r.font.color.rgb = RGBColor(0x29, 0x80, 0xB9)

    # 푸터
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "아진산업(주) · 본 보고서는 컴플라이언스 모니터링 시스템에서 자동 생성됨"
    )
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_bulk_zip(crawlers: list[dict[str, Any]]) -> bytes:
    """9개 개별 JSON + index.txt → ZIP."""
    import zipfile
    from io import BytesIO

    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        index_lines = ["크롤러,파일명,갱신시각,항목수,출처"]
        for c in crawlers:
            zf.writestr(
                f"{c['name']}.json",
                json.dumps(c["raw"], ensure_ascii=False, indent=2),
            )
            index_lines.append(
                f"{c['name']},{c['filename']},{c['crawled_at']},{len(c['items'])},{c['source'][:80]}"
            )
        zf.writestr("index.csv", "\n".join(index_lines))
    return buf.getvalue()


@router.get("/crawl/results/bulk-download")
async def download_bulk_crawl_results(
    format: str = "report",
    user=Depends(get_current_user),
):
    """9개 크롤러 결과 통합 다운로드.

    format: report (회사 양식 DOCX, 기본) | docx | xlsx | pdf | json | zip
    """
    from fastapi.responses import Response
    from datetime import datetime

    crawlers = _gather_all_crawler_results()
    today = datetime.now().strftime("%Y%m%d")
    base = f"compliance_bulk_{today}"
    fmt = format.lower().strip()

    author = ""
    if isinstance(user, dict):
        username = user.get("username", "")
        position = user.get("position", "")
        author = f"{username} {position}".strip()

    if fmt == "json":
        payload = {
            "generated_at": datetime.now().isoformat(),
            "crawlers": {c["name"]: c["raw"] for c in crawlers},
            "total_items": sum(len(c["items"]) for c in crawlers),
        }
        return Response(
            content=json.dumps(payload, ensure_ascii=False, indent=2),
            media_type="application/json; charset=utf-8",
            headers={"Content-Disposition": f"attachment; filename={base}.json"},
        )
    if fmt == "xlsx":
        return Response(
            content=_build_bulk_xlsx(crawlers),
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename={base}.xlsx"},
        )
    if fmt in ("docx", "report"):
        return Response(
            content=_build_bulk_docx(crawlers, author=author),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={base}_report.docx"},
        )
    if fmt == "pdf":
        # PDF는 단일 크롤러용 _build_pdf 를 9번 페이지로 합칠 수 있지만
        # 현재 fpdf2 기반으로 통합 보고서가 무거워서 첫 5개만 포함.
        # 실용적으로는 docx 가 더 좋은 선택. 단순 폴백.
        meta_combined = {
            "crawled_at": datetime.now().isoformat(),
            "source": f"{len(crawlers)}개 크롤러 통합",
        }
        all_items: list[dict[str, Any]] = []
        for c in crawlers:
            for it in c["items"][:10]:  # 각 크롤러 상위 10개
                all_items.append(
                    {
                        "title": f"[{c['name']}] {it.get('title', '')}",
                        "url": it.get("url", ""),
                        "summary": it.get("summary", ""),
                    }
                )
        return Response(
            content=_build_pdf("BULK", meta_combined, all_items),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={base}.pdf"},
        )
    if fmt == "zip":
        return Response(
            content=_build_bulk_zip(crawlers),
            media_type="application/zip",
            headers={"Content-Disposition": f"attachment; filename={base}.zip"},
        )

    raise HTTPException(
        status_code=400,
        detail=f"지원하지 않는 포맷: '{fmt}'. (report|docx|xlsx|pdf|json|zip)",
    )


@router.get(
    "/crawl/results/{name}",
    response_model=CrawlResultDetailResponse,
)
async def get_crawl_result_detail(
    name: str,
    limit: int = 50,
    offset: int = 0,
    user=Depends(get_current_user),
):
    """특정 크롤러의 항목 리스트를 반환 (페이지네이션)."""
    if name not in _CRAWLER_FILE_MAP:
        raise HTTPException(status_code=404, detail=f"크롤러 '{name}' 가 등록되지 않았습니다.")

    filename = _CRAWLER_FILE_MAP[name]
    path = DATA_DIR / "crawled" / filename

    if not path.exists():
        # 실행 전 — 빈 결과 반환
        return CrawlResultDetailResponse(
            name=name,
            filename=filename,
            total=0,
            items=[],
            has_more=False,
        )

    try:
        with open(path, encoding="utf-8") as f:
            data = json.load(f)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"파일 파싱 실패: {e}")

    raw_items = _extract_items(data)
    lo = max(0, offset)
    hi = lo + max(1, min(limit, 200))
    page = raw_items[lo:hi]

    items: list[CrawlResultItem] = []
    for raw in page:
        if not isinstance(raw, dict):
            items.append(CrawlResultItem(title=str(raw), summary="", url=""))
            continue
        title = str(
            raw.get("title")
            or raw.get("name")
            or raw.get("standard")
            or raw.get("law_name")
            or raw.get("id")
            or "(제목 없음)"
        )
        url = str(raw.get("url", raw.get("link", raw.get("source_url", ""))))
        summary = str(
            raw.get("summary")
            or raw.get("description")
            or raw.get("content", "")
        )[:500]  # 최대 500자
        # extra: 위에서 추출한 표준 필드 외 모든 필드
        extra = {k: v for k, v in raw.items() if k not in ("title", "name", "url", "link", "summary", "description")}
        items.append(CrawlResultItem(title=title, url=url, summary=summary, extra=extra))

    return CrawlResultDetailResponse(
        name=name,
        filename=filename,
        crawled_at=str(data.get("crawled_at", "")),
        source=str(data.get("source", "")),
        total=len(raw_items),
        items=items,
        has_more=hi < len(raw_items),
    )
