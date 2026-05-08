"""
Jinja2 템플릿 → DOCX / PDF / HWPX 변환 모듈

두 가지 모드 지원:
  1. 빈 양식 (blank): {{ 변수 }} → [___변수명___] 공란으로 렌더링
  2. 채운 양식 (filled): 사용자 입력값으로 렌더링

변환 결과는 바이트 버퍼로 반환 (st.download_button에 직접 전달 가능)
"""

import io
import re
import tempfile
from pathlib import Path

from features.draft.template_catalog import TEMPLATE_DIR, find_template_by_id


def _get_jinja_env():
    """Jinja2 환경 생성 (보안: SandboxedEnvironment)"""
    from jinja2 import Environment, FileSystemLoader
    try:
        from jinja2 import SandboxedEnvironment
        env = SandboxedEnvironment(
            loader=FileSystemLoader(str(TEMPLATE_DIR)),
            autoescape=False,
            keep_trailing_newline=True,
        )
    except ImportError:
        env = Environment(
            loader=FileSystemLoader(str(TEMPLATE_DIR)),
            autoescape=False,
        )
    return env


def render_template_blank(template_filename: str) -> str:
    """빈 양식 렌더링: 모든 변수를 [___변수명___] 공란으로 대체"""
    template_path = TEMPLATE_DIR / template_filename
    if not template_path.exists():
        return f"[오류] 템플릿 파일을 찾을 수 없습니다: {template_filename}"

    raw = template_path.read_text(encoding="utf-8")

    def replace_var(match):
        var_name = match.group(1).strip()
        display = var_name.replace("_", " ").title()
        return f"[___ {display} ___]"

    result = re.sub(r"\{\{[\s]*([^}]+?)[\s]*\}\}", replace_var, raw)
    result = re.sub(
        r"\{%[\s]*(?:if|elif|else|endif|for|endfor|block|endblock|extends|include|macro|endmacro).*?%\}",
        "", result,
    )

    return result.strip()


def render_template_filled(template_filename: str, values: dict) -> str:
    """채운 양식 렌더링: 사용자 입력값으로 Jinja2 렌더링

    v3.1: catalog field key → Jinja2 변수명 자동 매핑 + 템플릿 변수 자동 감지
    """
    env = _get_jinja_env()

    try:
        template = env.get_template(template_filename)

        # v3.1: catalog field key → Jinja2 변수명 매핑
        mapped = _map_field_keys_to_template_vars(template_filename, values)

        # 리스트 키 자동 변환 (줄바꿈 → 리스트) — 본문/설명 필드는 제외
        _text_fields = {"main_body", "opening_paragraph", "closing_paragraph", "body",
                        "d2_problem", "defect_summary", "description", "content"}
        safe_values = {}
        for k, v in mapped.items():
            if isinstance(v, str) and "\n" in v and k not in _text_fields:
                safe_values[k] = [line.strip() for line in v.split("\n") if line.strip()]
            else:
                safe_values[k] = v

        return template.render(**safe_values).strip()
    except Exception as e:
        return f"[렌더링 오류] {e}"


# ── v3.1: Catalog field key → Jinja2 변수명 매핑 ──

# 공통 매핑: catalog의 간략한 key → 여러 가능한 Jinja2 변수명
_FIELD_KEY_EXPANSIONS = {
    "recipient": ["recipient_name", "recipient_department", "recipient_title", "customer_name"],
    "sender": ["sender_name", "sender_department", "sender_title"],
    "body": ["main_body", "opening_paragraph", "closing_paragraph"],
    "cc": ["cc_list"],
    "subject": ["subject"],
    "part_name": ["part_name"],
    "part_number": ["part_number"],
    "date": ["date", "created_date", "claim_date"],
    "doc_number": ["doc_number"],
    "author": ["author"],
    "department": ["department"],
}


def _map_field_keys_to_template_vars(template_filename: str, values: dict) -> dict:
    """
    사용자 폼 입력(간략 key) → Jinja2 템플릿 변수(상세 key)로 확장

    전략:
    1. 원본 values를 그대로 보존
    2. 각 값이 매핑 가능한 Jinja2 변수명에 추가 복사
    3. 'recipient' → recipient_name, recipient_department 등으로 분해
    """
    result = dict(values)  # 원본 보존

    # 'recipient' 스마트 분해: "품질관리팀 홍길동 과장님" → 부서/이름/직급
    if "recipient" in values and values["recipient"]:
        _expand_person_field(result, values["recipient"], "recipient")

    # 'sender' 스마트 분해
    if "sender" in values and values["sender"]:
        _expand_person_field(result, values["sender"], "sender")

    # 'body' → main_body (문자열 유지, 리스트 변환 안 함)
    if "body" in values and values["body"]:
        body = values["body"]
        result.setdefault("main_body", body)
        result.setdefault("opening_paragraph", body)
        result.setdefault("closing_paragraph", "")

    # 'cc' → cc_list (쉼표 구분 리스트)
    if "cc" in values and values["cc"]:
        cc_val = values["cc"]
        if isinstance(cc_val, str):
            result["cc_list"] = [c.strip() for c in cc_val.replace("/", ",").split(",") if c.strip()]
        elif isinstance(cc_val, list):
            result["cc_list"] = cc_val

    # subject, part_name, part_number 등은 key 그대로 사용 (이미 result에 있음)

    # 날짜 자동 채움
    if "date" not in result:
        from datetime import date
        result["date"] = date.today().isoformat()
    result.setdefault("created_date", result.get("date", ""))

    # customer_name (OEM용)
    if "customer_name" not in result and "recipient" in values:
        result["customer_name"] = result.get("recipient_department", values["recipient"])

    # ── v3.3: 카탈로그 key → Jinja2 변수명 직접 매핑 (불일치 해소) ──

    # 8D Report
    if "defect_desc" in values:
        result.setdefault("defect_summary", values["defect_desc"])
        result.setdefault("d2_problem", values["defect_desc"])
    if "quantity" in values:
        result.setdefault("defect_quantity", values["quantity"])
    if "emergency" in values:
        result.setdefault("d3_containment", values["emergency"])
    if "customer" in values:
        result.setdefault("customer_name", values["customer"])
    if "plant_line" in values:
        result.setdefault("plant_line", values["plant_line"])

    # ECN 변경통보
    if "before_spec" in values:
        result.setdefault("before_description", values["before_spec"])
    if "after_spec" in values:
        result.setdefault("after_description", values["after_spec"])
    if "reason" in values:
        result.setdefault("change_reason", values["reason"])
    if "effective_date" in values:
        result.setdefault("issue_date", values["effective_date"])
    if "ecn_number" in values:
        result.setdefault("doc_number", values["ecn_number"])
    if "impact_scope" in values:
        result.setdefault("impact_scope", values["impact_scope"])

    # 회의록
    if "date" in values:
        result.setdefault("meeting_date", values["date"])
    if "location" in values:
        result.setdefault("meeting_place", values["location"])
    # attendees: 줄바꿈 텍스트 → 리스트 of dict 변환
    if "attendees" in values and isinstance(values["attendees"], str):
        _attendee_list = []
        for line in values["attendees"].split("\n"):
            name = line.strip().rstrip(",")
            if name:
                parts = name.replace("(", " ").replace(")", "").split()
                if len(parts) >= 2:
                    _attendee_list.append({"name": parts[0], "department": parts[1], "title": parts[2] if len(parts) > 2 else ""})
                else:
                    _attendee_list.append({"name": name, "department": "", "title": ""})
        if _attendee_list:
            result["attendees"] = _attendee_list
    # agenda: 줄바꿈 텍스트 → 리스트 of dict 변환
    if "agenda" in values and isinstance(values["agenda"], str):
        _agenda_list = []
        for line in values["agenda"].split("\n"):
            item = line.strip().lstrip("0123456789.·- ")
            if item:
                _agenda_list.append({"title": item, "content": ""})
        if _agenda_list:
            result["agenda_items"] = _agenda_list
    # decisions: 줄바꿈 텍스트 → 리스트 of dict 변환
    if "decisions" in values and isinstance(values["decisions"], str):
        _decision_list = []
        for line in values["decisions"].split("\n"):
            item = line.strip().lstrip("0123456789.·- ")
            if item:
                _decision_list.append({"content": item, "owner": "", "deadline": ""})
        if _decision_list:
            result["decisions"] = _decision_list
    # action_items: 줄바꿈 텍스트 → 리스트 변환
    if "action_items" in values and isinstance(values["action_items"], str):
        _action_list = [line.strip() for line in values["action_items"].split("\n") if line.strip()]
        if _action_list:
            result["action_items"] = _action_list

    # 품질문제 개선대책서
    if "defect_description" in values:
        result.setdefault("defect_description", values["defect_description"])
    if "defect_quantity" in values:
        result.setdefault("defect_quantity", values["defect_quantity"])
    if "occurrence_date" in values:
        result.setdefault("occurrence_date", values["occurrence_date"])

    # 안전 인시던트 리포트
    if "incident_date" in values:
        result.setdefault("incident_date", values["incident_date"])
    if "incident_location" in values:
        result.setdefault("incident_location", values["incident_location"])
    if "incident_description" in values:
        result.setdefault("incident_description", values["incident_description"])

    return result


def _expand_person_field(result: dict, value: str, prefix: str):
    """'품질관리팀 홍길동 과장님' → prefix_department, prefix_name, prefix_title 분해"""
    parts = value.strip().split()

    # 부서명 접미사 패턴 (팀, 부, 원, 실, 처, 본부, Inc, Dept 등)
    _dept_suffixes = ("팀", "부", "원", "실", "처", "본부", "센터",
                      "Inc", "Inc.", "Dept", "Dept.", "Team", "HQ")

    if len(parts) >= 3:
        # 부서 이름 직급
        result.setdefault(f"{prefix}_department", parts[0])
        result.setdefault(f"{prefix}_name", parts[1])
        result.setdefault(f"{prefix}_title", " ".join(parts[2:]))
    elif len(parts) == 2:
        # v3.3: 첫 단어가 부서명 접미사로 끝나면 "부서 이름"으로 파싱
        if any(parts[0].endswith(s) for s in _dept_suffixes):
            result.setdefault(f"{prefix}_department", parts[0])
            result.setdefault(f"{prefix}_name", parts[1])
            result.setdefault(f"{prefix}_title", "")
        else:
            # 이름 직급
            result.setdefault(f"{prefix}_name", parts[0])
            result.setdefault(f"{prefix}_title", parts[1])
            result.setdefault(f"{prefix}_department", "")
    elif len(parts) == 1:
        result.setdefault(f"{prefix}_name", parts[0])
        result.setdefault(f"{prefix}_department", "")
        result.setdefault(f"{prefix}_title", "")

    # customer_name (OEM용)
    if prefix == "recipient" and "customer_name" not in result:
        result["customer_name"] = result.get("recipient_department", value)


def text_to_docx_bytes(text: str, title: str = "AJIN 문서") -> bytes:
    """텍스트 → DOCX 바이트 변환"""
    try:
        from features.draft.docx_exporter import DocxExporter
        exporter = DocxExporter()
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = Path(tmp.name)
        exporter.export(markdown_text=text, output_path=tmp_path, doc_title=title, doc_type="")
        data = tmp_path.read_bytes()
        tmp_path.unlink(missing_ok=True)
        return data
    except Exception:
        pass

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for line in text.split("\n"):
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
                continue
            if stripped.startswith(("■", "▪", "●", "·", "-")):
                doc.add_paragraph(stripped, style="List Bullet")
            elif stripped.startswith(("D1", "D2", "D3", "D4", "D5", "D6", "D7", "D8")):
                p = doc.add_paragraph()
                run = p.add_run(stripped)
                run.bold = True
            else:
                doc.add_paragraph(stripped)

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
    except ImportError:
        return b""


def text_to_pdf_bytes(text: str, title: str = "AJIN 문서") -> bytes:
    """텍스트 → PDF 바이트 변환"""
    try:
        from features.draft.pdf_exporter import PdfExporter
        exporter = PdfExporter()
        return exporter.export_bytes(markdown_text=text, doc_title=title)
    except Exception:
        return b""


def text_to_hwpx_bytes(text: str, title: str = "AJIN 문서") -> bytes:
    """텍스트 → HWPX 바이트 변환"""
    try:
        from features.draft.hwpx_exporter import HwpxExporter
        exporter = HwpxExporter()
        return exporter.export_bytes(markdown_text=text, doc_title=title, doc_type="")
    except Exception:
        return b""


def export_template(
    template_id: str,
    output_format: str,
    values: dict | None = None,
) -> tuple[bytes, str]:
    """양식 내보내기 통합 함수"""
    meta = find_template_by_id(template_id)
    if not meta:
        return b"", "error.txt"

    filename_base = meta["file"].replace(".j2", "")
    title = meta["name"]

    if values:
        text = render_template_filled(meta["file"], values)
        suffix = "filled"
    else:
        text = render_template_blank(meta["file"])
        suffix = "blank"

    # v3.5: CSV/XLSX 추가
    from features.draft.tabular_exporter import text_to_csv_bytes, text_to_xlsx_bytes

    ext_map = {
        "docx": (".docx", text_to_docx_bytes),
        "pdf":  (".pdf",  text_to_pdf_bytes),
        # v3.3: .hwpx → .odt (내부가 ODT 형식이므로 확장자 일치시켜 크로스플랫폼 호환)
        "hwpx": (".odt", text_to_hwpx_bytes),
        "csv":  (".csv", text_to_csv_bytes),
        "xlsx": (".xlsx", text_to_xlsx_bytes),
    }

    ext, converter = ext_map.get(output_format, (".txt", lambda t, n: t.encode("utf-8")))
    file_bytes = converter(text, title)
    out_filename = f"AJIN_{filename_base}_{suffix}{ext}"

    return file_bytes, out_filename
