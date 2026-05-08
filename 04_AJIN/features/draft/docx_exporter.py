"""Phase 5: .docx 출력 모듈 (고도화)

렌더링된 마크다운 텍스트를 python-docx를 사용하여 .docx 파일로 변환한다.
- 아진산업 로고 헤더
- 결재란 (작성 → 검토 → 승인)
- 페이지 번호 푸터
- 마크다운 파싱 (표, 목록, 제목, 볼드)
"""
from __future__ import annotations

import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn


# 아진산업 브랜드 색상
_BRAND_NAVY = RGBColor(0, 51, 102)
_BRAND_GOLD = RGBColor(180, 140, 50)
_GRAY = RGBColor(100, 100, 100)
_LIGHT_GRAY = RGBColor(200, 200, 200)
_WHITE = RGBColor(255, 255, 255)

_FONT_KR = "맑은 고딕"
_FONT_EN = "Calibri"

# ──────────────────────────────────────────
# 문서 유형별 레이아웃 설정
# ──────────────────────────────────────────

_DEFAULT_LAYOUT = {
    "show_title_block": True,
    "show_approval": True,
    "meta_block": None,
}

_DOC_LAYOUTS = {
    # ── 이메일: 결재란 없음, 수신/발신 메타 블록 ──
    "email_oem": {
        "show_title_block": True,
        "show_approval": False,
        "meta_block": "email",
    },
    "email_supplier": {
        "show_title_block": True,
        "show_approval": False,
        "meta_block": "email",
    },
    "email_internal": {
        "show_title_block": True,
        "show_approval": False,
        "meta_block": "email",
    },
    "email_overseas": {
        "show_title_block": True,
        "show_approval": False,
        "meta_block": "email",
    },
    # ── 8D 보고서: 결재란 + 8D 헤더 메타 블록 ──
    "report_8d": {
        "show_title_block": True,
        "show_approval": True,
        "meta_block": "report_8d",
    },
    # ── ECN: 결재란 + ECN 메타 블록 (변경 전/후) ──
    "report_ecn": {
        "show_title_block": True,
        "show_approval": True,
        "meta_block": "report_ecn",
    },
    # ── 회의록: 결재란 없음, 회의 메타 블록 ──
    "report_meeting": {
        "show_title_block": True,
        "show_approval": False,
        "meta_block": "meeting",
    },
}


class DocxExporter:
    """마크다운 초안을 .docx 파일로 변환한다."""

    def __init__(self):
        self._logo_path = Path(__file__).parent.parent.parent / "ui" / "assets" / "ajin_logo.svg"

    def export(
        self,
        markdown_text: str,
        output_path: Path,
        doc_title: str = "",
        include_approval: bool = True,
        author: str = "",
        doc_type: str = "",
    ) -> Path:
        """마크다운 텍스트를 .docx로 변환하여 저장한다.

        Args:
            doc_type: 문서 유형 (email_oem, report_8d, report_ecn, report_meeting 등)
                      유형에 따라 레이아웃이 자동 결정된다.
        """
        doc = Document()
        layout = _DOC_LAYOUTS.get(doc_type, _DEFAULT_LAYOUT)

        self._setup_page(doc)
        self._setup_styles(doc)
        self._add_header_footer(doc)

        # 유형별 상단 구성
        if layout.get("show_title_block", True):
            self._add_title_block(doc, doc_title)

        if layout.get("show_approval", include_approval):
            self._add_approval_table(doc, author)

        # 유형별 메타 정보 블록
        if layout.get("meta_block"):
            self._add_meta_block(doc, layout["meta_block"], markdown_text)

        doc.add_paragraph()  # 간격

        self._parse_markdown(doc, markdown_text)

        # 하단 회사 정보
        self._add_company_footer(doc)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path

    def export_bytes(
        self,
        markdown_text: str,
        doc_type: str = "",
        doc_title: str = "",
        include_approval: bool = True,
        author: str = "",
    ) -> bytes:
        """파일 시스템 거치지 않고 .docx 바이트 직접 반환 (Cloud Run 라우터용)."""
        from io import BytesIO

        doc = Document()
        layout = _DOC_LAYOUTS.get(doc_type, _DEFAULT_LAYOUT)

        self._setup_page(doc)
        self._setup_styles(doc)
        self._add_header_footer(doc)
        if layout.get("show_title_block", True):
            self._add_title_block(doc, doc_title)
        if layout.get("show_approval", include_approval):
            self._add_approval_table(doc, author)
        if layout.get("meta_block"):
            self._add_meta_block(doc, layout["meta_block"], markdown_text)
        doc.add_paragraph()
        self._parse_markdown(doc, markdown_text)
        self._add_company_footer(doc)

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ──────────────────────────────────────────
    # 페이지 설정
    # ──────────────────────────────────────────

    def _setup_page(self, doc: Document):
        """A4 용지, 여백 설정"""
        section = doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    def _setup_styles(self, doc: Document):
        """문서 기본 스타일을 설정한다."""
        style = doc.styles["Normal"]
        font = style.font
        font.name = _FONT_KR
        font.size = Pt(10)
        pf = style.paragraph_format
        pf.space_after = Pt(4)
        pf.line_spacing = Pt(16)

        # 한글 폰트 fallback 설정
        rpr = style.element.get_or_add_rPr()
        rFonts = rpr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = doc.element.makeelement(qn("w:rFonts"), {})
            rpr.append(rFonts)
        rFonts.set(qn("w:eastAsia"), _FONT_KR)

    # ──────────────────────────────────────────
    # 헤더 / 푸터
    # ──────────────────────────────────────────

    def _add_header_footer(self, doc: Document):
        """페이지 헤더(회사명)와 푸터(페이지 번호)를 추가한다."""
        section = doc.sections[0]

        # ── 헤더: 아진산업(주) ──
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run("아진산업(주)  AJIN INDUSTRY CO., LTD.")
        run.font.size = Pt(8)
        run.font.color.rgb = _GRAY
        run.font.name = _FONT_KR

        # ── 푸터: 페이지 번호 ──
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 페이지 번호 필드
        run = fp.add_run()
        run.font.size = Pt(8)
        run.font.color.rgb = _GRAY
        fld_xml = (
            '<w:fldSimple xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' w:instr=" PAGE "><w:r><w:t>1</w:t></w:r></w:fldSimple>'
        )
        from lxml import etree
        fld_elem = etree.fromstring(fld_xml)
        run._r.append(fld_elem)

        run2 = fp.add_run(" / ")
        run2.font.size = Pt(8)
        run2.font.color.rgb = _GRAY

        run3 = fp.add_run()
        run3.font.size = Pt(8)
        run3.font.color.rgb = _GRAY
        fld_xml2 = (
            '<w:fldSimple xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' w:instr=" NUMPAGES "><w:r><w:t>1</w:t></w:r></w:fldSimple>'
        )
        fld_elem2 = etree.fromstring(fld_xml2)
        run3._r.append(fld_elem2)

    # ──────────────────────────────────────────
    # 타이틀 블록
    # ──────────────────────────────────────────

    def _add_title_block(self, doc: Document, title: str):
        """문서 상단 타이틀 블록"""
        # 회사명
        p_company = doc.add_paragraph()
        p_company.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_company.add_run("아진산업(주)")
        run.bold = True
        run.font.size = Pt(18)
        run.font.color.rgb = _BRAND_NAVY

        # 문서 제목
        if title:
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p_title.add_run(title)
            run.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(30, 30, 30)

        # 구분선
        p_line = doc.add_paragraph()
        p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_line.add_run("━" * 50)
        run.font.size = Pt(8)
        run.font.color.rgb = _BRAND_GOLD

    # ──────────────────────────────────────────
    # 결재란
    # ──────────────────────────────────────────

    def _add_approval_table(self, doc: Document, author: str = ""):
        """결재란 테이블 (작성 → 검토 → 승인)"""
        table = doc.add_table(rows=3, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.RIGHT

        # 열 너비 설정
        for row in table.rows:
            for i, width in enumerate([Cm(2.0), Cm(3.0), Cm(3.0), Cm(3.0)]):
                row.cells[i].width = width

        # 헤더 행
        headers = ["구분", "작성", "검토", "승인"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = _WHITE
            # 배경색 (navy)
            shading = cell._element.get_or_add_tcPr()
            shading_elem = doc.element.makeelement(qn("w:shd"), {
                qn("w:fill"): "003366",
                qn("w:val"): "clear",
            })
            shading.append(shading_elem)

        # 서명란 (빈 공간)
        sign_label = ["서명", "", "", ""]
        for i, label in enumerate(sign_label):
            cell = table.rows[1].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i == 0:
                run = p.add_run(label)
                run.font.size = Pt(8)
                run.font.color.rgb = _GRAY
            else:
                # 서명 공간 (높이 확보)
                p.add_run("\n\n")

        # 일자
        today = datetime.now().strftime("%Y.%m.%d")
        date_values = ["일자", today, "", ""]
        for i, val in enumerate(date_values):
            cell = table.rows[2].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.size = Pt(8)
            run.font.color.rgb = _GRAY

    # ──────────────────────────────────────────
    # 하단 회사 정보
    # ──────────────────────────────────────────

    def _add_company_footer(self, doc: Document):
        """문서 하단 회사 정보 블록"""
        doc.add_paragraph()  # 간격

        p_line = doc.add_paragraph()
        run = p_line.add_run("─" * 60)
        run.font.size = Pt(7)
        run.font.color.rgb = _LIGHT_GRAY

        info_text = (
            "아진산업(주)  |  경북 경산시 진량읍 공단8로 26길 40  |  TEL 053-856-9100\n"
            "AJIN INDUSTRY CO., LTD.  |  Confidential  |  본 문서는 사내 업무용으로 작성되었습니다."
        )
        p_info = doc.add_paragraph()
        p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_info.add_run(info_text)
        run.font.size = Pt(7)
        run.font.color.rgb = _GRAY

    # ──────────────────────────────────────────
    # 문서 유형별 메타 블록
    # ──────────────────────────────────────────

    def _add_meta_block(self, doc: Document, block_type: str, content: str):
        """문서 유형별 상단 메타 정보 블록"""
        today = datetime.now().strftime("%Y년 %m월 %d일")

        if block_type == "email":
            self._add_email_meta(doc, content, today)
        elif block_type == "report_8d":
            self._add_8d_meta(doc, content, today)
        elif block_type == "report_ecn":
            self._add_ecn_meta(doc, content, today)
        elif block_type == "meeting":
            self._add_meeting_meta(doc, content, today)

    def _add_email_meta(self, doc: Document, content: str, today: str):
        """이메일 수신/발신 메타 블록"""
        # 본문에서 수신/발신/제목 추출
        meta = {"수신": "", "발신": "아진산업 품질관리팀", "제목": "", "일자": today}
        for line in content.split("\n")[:10]:
            for key in meta:
                if line.startswith(f"{key}:") or line.startswith(f"{key} :"):
                    meta[key] = line.split(":", 1)[1].strip()

        table = doc.add_table(rows=4, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        for row in table.rows:
            row.cells[0].width = Cm(3.0)
            row.cells[1].width = Cm(13.0)

        items = [("수신", meta["수신"]), ("발신", meta["발신"]),
                 ("제목", meta["제목"]), ("일자", meta["일자"])]
        for r_idx, (label, value) in enumerate(items):
            # 라벨 셀
            cell_l = table.rows[r_idx].cells[0]
            cell_l.text = ""
            p = cell_l.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(label)
            run.bold = True
            run.font.size = Pt(9)
            # 배경색
            shading = cell_l._element.get_or_add_tcPr()
            shading_elem = doc.element.makeelement(qn("w:shd"), {
                qn("w:fill"): "E8EEF4", qn("w:val"): "clear",
            })
            shading.append(shading_elem)

            # 값 셀
            cell_v = table.rows[r_idx].cells[1]
            cell_v.text = ""
            p = cell_v.paragraphs[0]
            run = p.add_run(value)
            run.font.size = Pt(9)

    def _add_8d_meta(self, doc: Document, content: str, today: str):
        """8D 보고서 메타 블록 (발행번호, 고객, 부품, 일자)"""
        table = doc.add_table(rows=2, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        items = [
            ("발행번호", "8D-XXXX-XXX"), ("발행일자", today),
            ("고객사", ""), ("관련 부품", ""),
        ]
        for r_idx in range(2):
            for c_idx in range(2):
                idx = r_idx * 2 + c_idx
                label, value = items[idx]
                label_cell = table.rows[r_idx].cells[c_idx * 2]
                value_cell = table.rows[r_idx].cells[c_idx * 2 + 1] if c_idx * 2 + 1 < 4 else label_cell

                label_cell.text = ""
                p = label_cell.paragraphs[0]
                run = p.add_run(label)
                run.bold = True
                run.font.size = Pt(9)
                shading = label_cell._element.get_or_add_tcPr()
                shading_elem = doc.element.makeelement(qn("w:shd"), {
                    qn("w:fill"): "FFF3E0", qn("w:val"): "clear",
                })
                shading.append(shading_elem)

                if value_cell != label_cell:
                    value_cell.text = value
                    for p in value_cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(9)

    def _add_ecn_meta(self, doc: Document, content: str, today: str):
        """ECN 변경통보 메타 블록"""
        table = doc.add_table(rows=3, cols=4)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        items = [
            ("ECN 번호", "ECN-XXXX-XXX"), ("발행일", today),
            ("변경 유형", ""), ("긴급도", ""),
            ("대상 부품", ""), ("적용 시점", ""),
        ]
        for r_idx in range(3):
            for c_idx in range(2):
                idx = r_idx * 2 + c_idx
                label, value = items[idx]
                label_cell = table.rows[r_idx].cells[c_idx * 2]
                value_cell = table.rows[r_idx].cells[c_idx * 2 + 1] if c_idx * 2 + 1 < 4 else label_cell

                label_cell.text = ""
                p = label_cell.paragraphs[0]
                run = p.add_run(label)
                run.bold = True
                run.font.size = Pt(9)
                shading = label_cell._element.get_or_add_tcPr()
                shading_elem = doc.element.makeelement(qn("w:shd"), {
                    qn("w:fill"): "E3F2FD", qn("w:val"): "clear",
                })
                shading.append(shading_elem)

                if value_cell != label_cell:
                    value_cell.text = value
                    for p in value_cell.paragraphs:
                        for run in p.runs:
                            run.font.size = Pt(9)

    def _add_meeting_meta(self, doc: Document, content: str, today: str):
        """회의록 메타 블록"""
        table = doc.add_table(rows=3, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        for row in table.rows:
            row.cells[0].width = Cm(3.0)
            row.cells[1].width = Cm(13.0)

        items = [("회의일시", today), ("회의장소", ""), ("참석자", "")]
        for r_idx, (label, value) in enumerate(items):
            cell_l = table.rows[r_idx].cells[0]
            cell_l.text = ""
            p = cell_l.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(label)
            run.bold = True
            run.font.size = Pt(9)
            shading = cell_l._element.get_or_add_tcPr()
            shading_elem = doc.element.makeelement(qn("w:shd"), {
                qn("w:fill"): "E8F5E9", qn("w:val"): "clear",
            })
            shading.append(shading_elem)

            cell_v = table.rows[r_idx].cells[1]
            cell_v.text = value
            for p in cell_v.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    # ──────────────────────────────────────────
    # 마크다운 파싱
    # ──────────────────────────────────────────

    def _parse_markdown(self, doc: Document, text: str):
        """마크다운 텍스트를 파싱하여 docx 요소로 변환한다."""
        lines = text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]

            # H1
            if line.startswith("# "):
                h = doc.add_heading(line[2:].strip(), level=1)
                for run in h.runs:
                    run.font.color.rgb = _BRAND_NAVY
                i += 1
                continue

            # H2
            if line.startswith("## "):
                h = doc.add_heading(line[3:].strip(), level=2)
                for run in h.runs:
                    run.font.color.rgb = _BRAND_NAVY
                i += 1
                continue

            # H3
            if line.startswith("### "):
                doc.add_heading(line[4:].strip(), level=3)
                i += 1
                continue

            # 표
            if line.strip().startswith("|"):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i])
                    i += 1
                self._add_table(doc, table_lines)
                continue

            # 구분선
            if line.strip() in ("---", "───", "─" * 50, "***"):
                p = doc.add_paragraph()
                run = p.add_run("─" * 50)
                run.font.size = Pt(8)
                run.font.color.rgb = _LIGHT_GRAY
                i += 1
                continue

            # 볼드 목록 (- **부서**: 내용)
            bold_match = re.match(r"^-\s+\*\*(.+?)\*\*:\s*(.+)$", line.strip())
            if bold_match:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(bold_match.group(1))
                run.bold = True
                p.add_run(f": {bold_match.group(2)}")
                i += 1
                continue

            # 목록 항목
            if line.strip().startswith("- ") or line.strip().startswith("· "):
                doc.add_paragraph(line.strip()[2:], style="List Bullet")
                i += 1
                continue

            # 번호 목록
            if re.match(r"^\s*\d+\.", line):
                doc.add_paragraph(
                    re.sub(r"^\s*\d+\.\s*", "", line), style="List Number"
                )
                i += 1
                continue

            # ■ 구조화 항목
            if line.strip().startswith("■"):
                p = doc.add_paragraph()
                run = p.add_run(line.strip())
                run.bold = True
                i += 1
                continue

            # 인라인 볼드 처리
            if line.strip() and "**" in line:
                p = doc.add_paragraph()
                self._add_inline_bold(p, line.strip())
                i += 1
                continue

            # 일반 텍스트
            if line.strip():
                doc.add_paragraph(line.strip())

            i += 1

    def _add_inline_bold(self, paragraph, text: str):
        """인라인 **볼드** 처리"""
        parts = re.split(r"(\*\*.*?\*\*)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)

    def _add_table(self, doc: Document, table_lines: list[str]):
        """마크다운 표를 docx 테이블로 변환한다."""
        data_lines = [
            line for line in table_lines if not re.match(r"^\|[\s\-:|]+\|$", line)
        ]
        if not data_lines:
            return

        rows = []
        for line in data_lines:
            cells = [c.strip() for c in line.strip("|").split("|")]
            rows.append(cells)

        if not rows:
            return

        num_cols = len(rows[0])
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for r_idx, row in enumerate(rows):
            for c_idx, cell_text in enumerate(row):
                if c_idx < num_cols:
                    cell = table.rows[r_idx].cells[c_idx]
                    cell.text = cell_text

                    # 헤더 행 스타일
                    if r_idx == 0:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                                run.font.size = Pt(9)
                        # 헤더 배경색
                        shading = cell._element.get_or_add_tcPr()
                        shading_elem = doc.element.makeelement(qn("w:shd"), {
                            qn("w:fill"): "E8EEF4",
                            qn("w:val"): "clear",
                        })
                        shading.append(shading_elem)
                    else:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(9)
